import fitz
import re
import os
from groq import Groq
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from utils.file_utils import create_zip
import streamlit as st


# ==============================
# CONFIG
# ==============================
MODEL_NAME = "llama-3.1-8b-instant"
CHUNK_SIZE = 3500
REDUCE_BATCH_SIZE = 3


# ==============================
# TEXT EXTRACTION
# ==============================
def extract_text_from_pdf_bytes(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    return full_text


def chunk_text(text, chunk_size=CHUNK_SIZE):
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]


# ==============================
# TITLE & AUTHOR EXTRACTION
# ==============================
def extract_title_and_authors_from_bytes(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[0]
    blocks = page.get_text("dict")["blocks"]

    rows = []
    for b in blocks:
        if "lines" not in b:
            continue

        text = " ".join(
            span["text"]
            for line in b["lines"]
            for span in line["spans"]
        ).strip()

        if not text:
            continue

        max_font = max(
            span["size"]
            for line in b["lines"]
            for span in line["spans"]
        )

        y0 = b["bbox"][1]
        rows.append((text, max_font, y0))

    rows.sort(key=lambda x: x[2])
    page_height = page.rect.height

    candidates = []

    for text, size, y in rows:
        if y > page_height * 0.45:
            break

        tl = text.lower()
        wc = len(text.split())

        if re.search(r"(received|accepted|published|submitted)", tl):
            continue
        if wc <= 5 and "&" in text:
            continue
        if wc <= 5 and text.istitle():
            continue
        if wc <= 6 and re.search(r"(journal|letters|review|transactions|proceedings)", tl):
            continue
        if tl.startswith("and "):
            continue
        if wc <= 2:
            continue

        candidates.append((text, size))

    if candidates:
        raw_title = sorted(candidates, key=lambda x: (-x[1], -len(x[0])))[0][0]
    else:
        raw_title = rows[0][0]

    title = re.sub(r"\s+", " ", raw_title).strip()

    authors = "Not explicitly detected"
    title_seen = False

    for text, _, _ in rows:
        if title in text:
            title_seen = True
            continue

        if not title_seen:
            continue

        clean = text.strip()
        cl = clean.lower()

        if re.search(r"(doi|abstract|keywords)", cl):
            continue
        if re.search(r"\b(19|20)\d{2}\b", clean):
            continue
        if clean.count(",") == 0 and " and " not in cl:
            continue

        if 5 <= len(clean) <= 200:
            authors = clean
            break

    return title, authors


# ==============================
# LLM CALLS
# ==============================
def summarize_chunk(client, chunk):
    prompt = f"""
You are analyzing a research paper.

Extract key technical insights from the following text.
Focus on:
- Problem Statement
- Methodology
- Models/Algorithms
- Datasets
- Evaluation Metrics
- Key Findings
- Limitations

Text:
{chunk}
"""

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
    )

    return response.choices[0].message.content


def reduce_notes_in_batches(client, notes, batch_size=REDUCE_BATCH_SIZE):
    reduced_batches = []

    for i in range(0, len(notes), batch_size):
        batch = "\n\n".join(notes[i:i + batch_size])

        prompt = f"""
Consolidate the following extracted notes into a structured technical summary.

Remove repetition.
Preserve important findings.
Keep it concise but complete.

Notes:
{batch}
"""

        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )

        reduced_batches.append(response.choices[0].message.content)

    if len(reduced_batches) == 1:
        return reduced_batches[0]

    return reduce_notes_in_batches(client, reduced_batches, batch_size)


def generate_one_pager(client, title, authors, reduced_notes):
    prompt = f"""
Generate a professional 1-page research summary in the following format:

Title: {title}
Authors: {authors}

1. Background
2. Objective
3. Methodology
4. Key Results
5. Strengths
6. Limitations
7. Future Scope
8. Reference (Include authors and year if mentioned in paper)

Content:
{reduced_notes}
"""

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
    )

    return response.choices[0].message.content


# ==============================
# WORD EXPORT (PROFESSIONAL STYLE)
# ==============================
def save_summary_to_word(summary_text, output_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    import re

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = summary_text.split("\n")

    for line in lines:
        stripped = line.strip()

        # Skip completely empty lines safely
        if not stripped:
            doc.add_paragraph("")  # keep spacing
            continue

        # TITLE
        if stripped.startswith("Title:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)

        # AUTHORS
        elif stripped.startswith("Authors:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.italic = True
            run.font.size = Pt(11)

        # SECTION HEADERS (1. Background etc.)
        elif re.match(r"^\d+\.", stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)

        # NORMAL TEXT
        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(11)

    doc.save(output_path)


# ==============================
# MAIN ENTRY FUNCTION (STREAMLIT CALL)
# ==============================
def summarize_pdfs(pdf_files, output_dir):
    """
    pdf_files: Dict[str, bytes]
    returns: Dict[str, bytes]  (docx files)
    """

    import streamlit as st
    from io import BytesIO
    import re

    try:
        api_key = st.secrets["GROQ_API_KEY"]
    except KeyError:
        raise ValueError("GROQ_API_KEY not found in Streamlit secrets")

    client = Groq(api_key=api_key)

    summaries_dict = {}

    total_pdfs = len(pdf_files)

    # 🔹 Overall progress
    overall_progress = st.progress(0)
    overall_status = st.empty()

    for pdf_index, (filename, pdf_bytes) in enumerate(pdf_files.items(), start=1):

        overall_status.markdown(
            f"### 📄 Processing **{filename}** ({pdf_index}/{total_pdfs})"
        )

        # 🔹 Per-PDF progress
        pdf_progress = st.progress(0)
        pdf_status = st.empty()

        # 1️⃣ Extract title & authors
        pdf_status.info("🔍 Extracting title & authors")
        title, authors = extract_title_and_authors_from_bytes(pdf_bytes)
        pdf_progress.progress(10)

        # 2️⃣ Extract text
        pdf_status.info("📄 Extracting text from PDF")
        text = extract_text_from_pdf_bytes(pdf_bytes)
        pdf_progress.progress(20)

        # 3️⃣ Chunk text
        chunks = chunk_text(text)
        total_chunks = len(chunks)

        pdf_status.info(f"🔪 Split into {total_chunks} chunks")
        pdf_progress.progress(30)

        # 4️⃣ Summarize chunks
        notes = []
        for i, chunk in enumerate(chunks, start=1):
            pdf_status.info(f"🧠 Summarizing chunk {i}/{total_chunks}")
            notes.append(summarize_chunk(client, chunk))

            # Progress from 30 → 70
            pdf_progress.progress(30 + int(40 * i / total_chunks))

        # 5️⃣ Reduce notes
        pdf_status.info("🧩 Consolidating notes")
        reduced = reduce_notes_in_batches(client, notes)
        pdf_progress.progress(80)

        # 6️⃣ Generate final one-pager
        pdf_status.info("📝 Generating 1-pager summary")
        final_summary = generate_one_pager(client, title, authors, reduced)
        pdf_progress.progress(90)

        # 7️⃣ Save Word file to memory
        pdf_status.info("💾 Creating Word document")

        safe_title = re.sub(r'[\\/*?:"<>|]', "", title)[:60]
        output_filename = f"{safe_title}.docx"

        buffer = BytesIO()
        save_summary_to_word(final_summary, buffer)

        summaries_dict[output_filename] = buffer.getvalue()

        pdf_progress.progress(100)
        pdf_status.success("✅ Completed")

        # 🔹 Update overall progress
        overall_progress.progress(int(100 * pdf_index / total_pdfs))

    overall_status.success("🎉 All PDFs summarized successfully!")

    return summaries_dict
